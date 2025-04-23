"""
Document Extractor Module

This module provides functionality to extract text from various document formats
including PDF, DOCX, and plain text files.
"""

import os
import logging
from typing import Dict, Any, Optional, List, Tuple
import PyPDF2
import docx
import textract
from pdfminer.high_level import extract_text as pdfminer_extract_text
import fitz  # PyMuPDF for PDF image extraction
from .image_utils import extract_images_from_pdf, extract_images_from_docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentExtractor:
    """
    A class for extracting text from various document formats.
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx, .doc)
    - Text (.txt)
    - Other formats supported by textract
    """
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_doc,
            '.txt': self._extract_from_txt,
        }
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata
                - pages: List of pages (for multi-page documents)
                - error: Error message if extraction failed
        """
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}"}
        
        _, ext = os.path.splitext(file_path.lower())
        
        if ext not in self.supported_extensions:
            return self._extract_with_textract(file_path)
        
        try:
            return self.supported_extensions[ext](file_path)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {"error": f"Extraction failed: {str(e)}"}
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files using PyPDF2 and pdfminer as fallback."""
        try:
            return self._extract_from_pdf_with_pypdf2(file_path)
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed, trying pdfminer: {str(e)}")
            return self._extract_from_pdf_with_pdfminer(file_path)
    
    def _extract_from_pdf_with_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files using PyPDF2."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Extract metadata
            metadata = reader.metadata
            if metadata:
                metadata_dict = {
                    'title': metadata.get('/Title', ''),
                    'author': metadata.get('/Author', ''),
                    'subject': metadata.get('/Subject', ''),
                    'creator': metadata.get('/Creator', ''),
                    'producer': metadata.get('/Producer', ''),
                }
            else:
                metadata_dict = {}
            
            # Extract text from each page
            pages = []
            full_text = ""
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append(text)
                    full_text += text + "\n\n"
            
            return {
                "text": full_text.strip(),
                "metadata": metadata_dict,
                "pages": pages,
                "page_count": len(reader.pages)
            }
    
    def _extract_from_pdf_with_pdfminer(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files using pdfminer."""
        full_text = pdfminer_extract_text(file_path)
        
        # Split text into pages (approximate)
        pages = full_text.split('\f')
        pages = [p.strip() for p in pages if p.strip()]
        
        return {
            "text": full_text.strip(),
            "metadata": {},  # pdfminer doesn't extract metadata easily
            "pages": pages,
            "page_count": len(pages)
        }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        doc = docx.Document(file_path)
        
        # Extract metadata
        core_properties = doc.core_properties
        metadata = {
            'title': core_properties.title or '',
            'author': core_properties.author or '',
            'subject': core_properties.subject or '',
            'created': str(core_properties.created) if core_properties.created else '',
            'modified': str(core_properties.modified) if core_properties.modified else '',
        }
        
        # Extract text from paragraphs
        paragraphs = []
        full_text = ""
        
        for para in doc.paragraphs:
            if para.text:
                paragraphs.append(para.text)
                full_text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                if row_text:
                    full_text += row_text + "\n"
        
        return {
            "text": full_text.strip(),
            "metadata": metadata,
            "paragraphs": paragraphs
        }
    
    def _extract_from_doc(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOC files using textract."""
        try:
            text = textract.process(file_path).decode('utf-8')
            return {
                "text": text.strip(),
                "metadata": {}
            }
        except Exception as e:
            logger.error(f"Error extracting text from DOC file: {str(e)}")
            return {"error": f"DOC extraction failed: {str(e)}"}
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Split into lines for basic structure
            lines = text.split('\n')
            
            return {
                "text": text.strip(),
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "size": os.path.getsize(file_path)
                },
                "lines": lines
            }
        except UnicodeDecodeError:
            # Try with different encoding if utf-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                lines = text.split('\n')
                
                return {
                    "text": text.strip(),
                    "metadata": {
                        "filename": os.path.basename(file_path),
                        "size": os.path.getsize(file_path),
                        "encoding": "latin-1"
                    },
                    "lines": lines
                }
            except Exception as e:
                return {"error": f"Text extraction failed: {str(e)}"}
    
    def _extract_with_textract(self, file_path: str) -> Dict[str, Any]:
        """Extract text using textract for unsupported file types."""
        try:
            text = textract.process(file_path).decode('utf-8')
            return {
                "text": text.strip(),
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Textract extraction failed: {str(e)}")
            return {"error": f"Unsupported file type extraction failed: {str(e)}"}
    
    def extract_with_images(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and images from a document file.
        Returns a dict with text, images, and metadata for Mistral 7B training data preparation.
        """
        base = self.extract(file_path)
        _, ext = os.path.splitext(file_path.lower())
        images = []
        if ext == ".pdf":
            images = extract_images_from_pdf(file_path)
        elif ext == ".docx":
            images = extract_images_from_docx(file_path)
        # Symbols: If symbols are images, they're included above; if text, handle in downstream processing.
        base["images"] = images
        return base


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Convenience function to extract text from a file.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    extractor = DocumentExtractor()
    return extractor.extract(file_path)
